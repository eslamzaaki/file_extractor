import pytest
import json
import tempfile
import os
from unittest.mock import Mock, patch
from app import app, extract_pdf, extract_docx, extract_csv, extract_txt, extract_doc

# Try to import libraries for creating test files
try:
    from docx import Document as DocxDocument
    DOCX_CREATE_AVAILABLE = True
except ImportError:
    DOCX_CREATE_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    PDF_CREATE_AVAILABLE = True
except ImportError:
    PDF_CREATE_AVAILABLE = False

@pytest.fixture
def client():
    """Create a test client for the Flask app"""
    app.config['TESTING'] = True
    with app.test_client() as client:
        yield client

@pytest.fixture
def sample_txt_file():
    """Create a real TXT file for testing"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write("Hello, this is a test text file.\nIt has multiple lines.\nFor testing purposes.")
        temp_path = f.name
    
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)

@pytest.fixture
def sample_csv_file():
    """Create a real CSV file for testing"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as f:
        f.write("Name,Age,City\nJohn,30,New York\nJane,25,Los Angeles\nBob,35,Chicago")
        temp_path = f.name
    
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)

@pytest.fixture
def sample_docx_file():
    """Create a real DOCX file for testing"""
    if not DOCX_CREATE_AVAILABLE:
        pytest.skip("python-docx not available for creating test files")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_path = f.name
    
    # Create a real DOCX file
    doc = DocxDocument()
    doc.add_paragraph("This is a test DOCX document.")
    doc.add_paragraph("It contains multiple paragraphs.")
    doc.add_paragraph("For testing the extraction functionality.")
    doc.save(temp_path)
    
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)

@pytest.fixture
def sample_pdf_file():
    """Create a real PDF file for testing"""
    if not PDF_CREATE_AVAILABLE:
        # Fallback: create minimal PDF manually
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            # Minimal valid PDF
            pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>
endobj
4 0 obj
<< /Length 65 >>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF Content for Extraction) Tj
0 -20 Td
(Second line of PDF text) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000256 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
380
%%EOF"""
            f.write(pdf_content)
            temp_path = f.name
    else:
        # Use reportlab to create a proper PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        c = canvas.Canvas(temp_path, pagesize=letter)
        c.drawString(100, 750, "Test PDF Content for Extraction")
        c.drawString(100, 730, "Second line of PDF text")
        c.drawString(100, 710, "This is a real PDF file created for testing.")
        c.save()
    
    yield temp_path
    if os.path.exists(temp_path):
        os.unlink(temp_path)


class TestExtractRoute:
    """Test cases for the /extract route - testing real extraction logic"""
    
    def test_extract_missing_url_get(self, client):
        """Test GET request without URL parameter"""
        response = client.get('/extract')
        assert response.status_code == 400
        data = json.loads(response.data)
        assert 'error' in data
        assert 'Missing file URL' in data['error']
    
    def test_extract_missing_url_post(self, client):
        """Test POST request without URL in body"""
        response = client.post('/extract', json={})
        assert response.status_code == 400
        data = json.loads(response.data)
        assert 'error' in data
        assert 'Missing file URL' in data['error']
    
    def test_extract_invalid_url(self, client):
        """Test with invalid URL that fails to download"""
        with patch('app.requests.get') as mock_get:
            mock_get.side_effect = Exception("Connection error")
            
            response = client.get('/extract?url=https://invalid-url.com/file.pdf')
            assert response.status_code == 400
            data = json.loads(response.data)
            assert 'error' in data
            assert 'Failed to download file' in data['error']
    
    @patch('app.requests.get')
    def test_extract_txt_file_get(self, mock_get, client, sample_txt_file):
        """Test extracting content from TXT file via GET - REAL EXTRACTION"""
        # Read the actual file content
        with open(sample_txt_file, 'rb') as f:
            file_content = f.read()
        
        # Mock only the HTTP download, use real file content
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'text/plain'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/test.txt')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert 'content' in data
        assert data['file_type'] == '.txt'
        assert 'Hello, this is a test text file' in data['content']
        assert 'multiple lines' in data['content']
        assert 'content_length' in data
        assert data['content_length'] > 0
    
    @patch('app.requests.get')
    def test_extract_txt_file_post(self, mock_get, client, sample_txt_file):
        """Test extracting content from TXT file via POST - REAL EXTRACTION"""
        with open(sample_txt_file, 'rb') as f:
            file_content = f.read()
        
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'text/plain'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.post('/extract', json={'url': 'https://example.com/test.txt'})
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert 'content' in data
        assert data['file_type'] == '.txt'
        assert 'For testing purposes' in data['content']
    
    @patch('app.requests.get')
    def test_extract_csv_file(self, mock_get, client, sample_csv_file):
        """Test extracting content from CSV file - REAL EXTRACTION"""
        with open(sample_csv_file, 'rb') as f:
            file_content = f.read()
        
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'text/csv'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/data.csv')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert 'content' in data
        assert data['file_type'] == '.csv'
        assert 'John' in data['content']
        assert 'Jane' in data['content']
        assert 'Bob' in data['content']
        assert 'New York' in data['content']
        assert 'Los Angeles' in data['content']
    
    @patch('app.requests.get')
    def test_extract_pdf_file(self, mock_get, client, sample_pdf_file):
        """Test extracting content from PDF file - REAL EXTRACTION"""
        with open(sample_pdf_file, 'rb') as f:
            file_content = f.read()
        
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'application/pdf'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/document.pdf')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert 'content' in data
        assert data['file_type'] == '.pdf'
        # PDF extraction may extract text differently, so check for any content
        assert len(data['content']) > 0
        assert 'content_length' in data
    
    @patch('app.requests.get')
    def test_extract_docx_file(self, mock_get, client, sample_docx_file):
        """Test extracting content from DOCX file - REAL EXTRACTION"""
        with open(sample_docx_file, 'rb') as f:
            file_content = f.read()
        
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/document.docx')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert 'content' in data
        assert data['file_type'] == '.docx'
        # Check for actual content from the DOCX file
        assert 'test DOCX document' in data['content'] or 'test' in data['content'].lower()
        assert len(data['content']) > 0
    
    @patch('app.requests.get')
    def test_extract_unsupported_file_type(self, mock_get, client):
        """Test with unsupported file type"""
        mock_response = Mock()
        mock_response.content = b"binary executable content"
        mock_response.headers = {'Content-Type': 'application/octet-stream'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/file.exe')
        assert response.status_code == 400
        data = json.loads(response.data)
        assert 'error' in data
        assert 'Unsupported file type' in data['error'] or 'failed to extract' in data['error'].lower()
    
    @patch('app.requests.get')
    def test_extract_file_detection_from_content_type(self, mock_get, client, sample_txt_file):
        """Test file type detection from Content-Type header"""
        with open(sample_txt_file, 'rb') as f:
            file_content = f.read()
        
        mock_response = Mock()
        mock_response.content = file_content
        mock_response.headers = {'Content-Type': 'text/plain'}
        mock_response.raise_for_status = Mock()
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/file')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['success'] is True
        assert data['file_type'] == '.txt'
        assert 'Hello' in data['content']
    
    @patch('app.requests.get')
    def test_extract_http_error(self, mock_get, client):
        """Test handling of HTTP errors (404, 500, etc.)"""
        mock_response = Mock()
        mock_response.raise_for_status.side_effect = Exception("404 Not Found")
        mock_get.return_value = mock_response
        
        response = client.get('/extract?url=https://example.com/missing.pdf')
        assert response.status_code == 400
        data = json.loads(response.data)
        assert 'error' in data
        assert 'Failed to download file' in data['error']
    
    @patch('app.requests.get')
    def test_extract_csv_with_special_characters(self, mock_get, client):
        """Test CSV extraction with special characters and encoding"""
        csv_content = "Name,Description,Price\nProduct 1,\"Description with, comma\",$10.99\nProduct 2,\"Multi\nline description\",$20.50"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as f:
            f.write(csv_content)
            temp_path = f.name
        
        try:
            with open(temp_path, 'rb') as f:
                file_content = f.read()
            
            mock_response = Mock()
            mock_response.content = file_content
            mock_response.headers = {'Content-Type': 'text/csv'}
            mock_response.raise_for_status = Mock()
            mock_get.return_value = mock_response
            
            response = client.get('/extract?url=https://example.com/special.csv')
            assert response.status_code == 200
            data = json.loads(response.data)
            assert data['success'] is True
            assert 'Product 1' in data['content']
            assert 'Product 2' in data['content']
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)


class TestExtractionFunctions:
    """Test cases for individual extraction functions - REAL EXTRACTION"""
    
    def test_extract_txt_function(self):
        """Test TXT extraction function with real file"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Test content\nLine 2\nLine 3\nSpecial chars: àáâãäå")
            temp_path = f.name
        
        try:
            content, error = extract_txt(temp_path)
            assert error is None
            assert content is not None
            assert "Test content" in content
            assert "Line 2" in content
            assert "Line 3" in content
        finally:
            os.unlink(temp_path)
    
    def test_extract_csv_function(self):
        """Test CSV extraction function with real file"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as f:
            f.write("Name,Age,Email\nJohn,30,john@example.com\nJane,25,jane@example.com")
            temp_path = f.name
        
        try:
            content, error = extract_csv(temp_path)
            assert error is None
            assert content is not None
            assert "John" in content
            assert "Jane" in content
            assert "john@example.com" in content
        finally:
            os.unlink(temp_path)
    
    def test_extract_txt_file_not_found(self):
        """Test TXT extraction with non-existent file"""
        content, error = extract_txt('/nonexistent/file.txt')
        assert content is None
        assert error is not None
    
    def test_extract_pdf_function(self, sample_pdf_file):
        """Test PDF extraction function with real PDF file"""
        content, error = extract_pdf(sample_pdf_file)
        # PDF extraction may or may not work depending on PDF structure
        # Just verify it doesn't crash and returns something
        if error is None:
            assert content is not None
        # If error, it's okay - some PDFs may not extract text properly
    
    def test_extract_docx_function(self, sample_docx_file):
        """Test DOCX extraction function with real DOCX file"""
        if not DOCX_CREATE_AVAILABLE:
            pytest.skip("python-docx not available")
        
        content, error = extract_docx(sample_docx_file)
        assert error is None
        assert content is not None
        assert len(content) > 0
        # Should contain some text from the document
        assert 'test' in content.lower() or 'document' in content.lower()


class TestOtherEndpoints:
    """Test cases for other endpoints"""
    
    def test_health_endpoint(self, client):
        """Test health check endpoint"""
        response = client.get('/health')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert data['status'] == 'healthy'
        assert 'pdf_support' in data
        assert 'docx_support' in data
        assert 'doc_support' in data
    
    def test_index_endpoint(self, client):
        """Test root endpoint"""
        response = client.get('/')
        assert response.status_code == 200
        data = json.loads(response.data)
        assert 'message' in data
        assert 'endpoints' in data
        assert 'supported_formats' in data
